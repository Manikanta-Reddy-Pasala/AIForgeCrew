"""Map-reduce document summariser — auto-selects strategy by size.

The LLM is monkeypatched so the tests are deterministic and offline: each fake
`complete` call echoes a short marker so we can assert map (per-window) and
reduce (fold) both ran.
"""
from __future__ import annotations

import docx

from aiforge_core.runtime import doc_summarize as ds


def _fake_complete(calls):
    """Return a `complete` stub that records every call and returns a marker
    encoding whether it was a MAP (per-window) or REDUCE (fold) prompt."""
    def _c(role, messages, *, max_tokens=None, **kw):
        prompt = messages[0]["content"]
        calls.append((role, max_tokens, prompt))
        if prompt.startswith("Below are section summaries"):
            return "FINAL-FOLD"
        return "map-sum"
    return _c


def test_small_text_returned_verbatim_no_llm(monkeypatch):
    """A tiny doc must NOT trigger an LLM call — cheap passthrough."""
    calls: list = []
    monkeypatch.setattr("aiforge_core.llm.client.complete", _fake_complete(calls))
    out = ds.summarize_text("short note about the deploy", role="chat")
    assert out == "short note about the deploy"
    assert calls == []


def test_large_text_maps_each_window_then_reduces(monkeypatch):
    calls: list = []
    monkeypatch.setattr("aiforge_core.llm.client.complete", _fake_complete(calls))
    monkeypatch.setenv("AIFORGE_SUMMARY_WINDOW_CHARS", "500")
    # ~10 windows of 500 chars.
    big = "\n".join(f"line {i} " + "x" * 90 for i in range(60))
    out = ds.summarize_text(big, role="chat")
    map_calls = [c for c in calls if not c[2].startswith("Below are section")]
    fold_calls = [c for c in calls if c[2].startswith("Below are section")]
    assert len(map_calls) >= 2          # each window summarised
    assert len(fold_calls) == 1         # one reduce fold
    assert out == "FINAL-FOLD"


def test_max_windows_cap_truncates(monkeypatch):
    calls: list = []
    monkeypatch.setattr("aiforge_core.llm.client.complete", _fake_complete(calls))
    monkeypatch.setenv("AIFORGE_SUMMARY_WINDOW_CHARS", "200")
    monkeypatch.setenv("AIFORGE_SUMMARY_MAX_WINDOWS", "3")
    big = "\n".join(f"line {i} " + "y" * 90 for i in range(100))
    out = ds.summarize_text(big, role="chat")
    map_calls = [c for c in calls if not c[2].startswith("Below are section")]
    assert len(map_calls) <= 3          # capped
    assert "exceeded 3 summary windows" in out


def test_disabled_falls_back_to_excerpt(monkeypatch):
    calls: list = []
    monkeypatch.setattr("aiforge_core.llm.client.complete", _fake_complete(calls))
    monkeypatch.setenv("AIFORGE_DOC_SUMMARY_ENABLED", "0")
    big = "z" * 20000
    out = ds.summarize_text(big, role="chat")
    assert calls == []                  # no LLM
    assert len(out) <= 6000             # excerpt fallback


def test_llm_failure_falls_back_to_excerpt(monkeypatch):
    def _boom(*a, **k):
        raise RuntimeError("endpoint down")
    monkeypatch.setattr("aiforge_core.llm.client.complete", _boom)
    monkeypatch.setenv("AIFORGE_SUMMARY_WINDOW_CHARS", "500")
    big = "\n".join(f"line {i} " + "w" * 90 for i in range(60))
    out = ds.summarize_text(big, role="chat")
    assert out                          # never empty
    assert len(out) <= 6000             # fell back to excerpt


def test_parse_page_spec():
    from aiforge_core.runtime import doc_extract as de
    assert de.parse_page_spec("10-12", 20) == [9, 10, 11]
    assert de.parse_page_spec("3,5,7-9", 20) == [2, 4, 6, 7, 8]
    assert de.parse_page_spec("12", 20) == [11]
    assert de.parse_page_spec("18-15", 20) == [14, 15, 16, 17]   # reversed range
    assert de.parse_page_spec("50", 20) == []                    # out of range
    assert de.parse_page_spec("junk", 20) == []
    assert de.parse_page_spec("", 20) == []


def test_docx_page_segmentation(tmp_path):
    """docx splits into pages on manual page breaks; each page keeps its text."""
    from aiforge_core.runtime import doc_extract as de
    d = docx.Document()
    d.add_paragraph("page one content")
    d.add_page_break()                      # manual break -> new page
    d.add_paragraph("page two content")
    d.add_page_break()
    d.add_paragraph("page three content")
    p = tmp_path / "multi.docx"
    d.save(str(p))
    pages = de.document_pages(str(p))
    assert len(pages) == 3
    assert "page one" in pages[0]
    assert "page two" in pages[1]
    assert "page three" in pages[2]


def test_extract_pages_selects_range(tmp_path):
    from aiforge_core.runtime import doc_extract as de
    d = docx.Document()
    for i in range(1, 6):
        d.add_paragraph(f"content of page {i}")
        if i < 5:
            d.add_page_break()
    p = tmp_path / "five.docx"
    d.save(str(p))
    out = de.extract_pages(str(p), "", "2-3")
    assert "--- page 2 ---" in out
    assert "--- page 3 ---" in out
    assert "content of page 2" in out
    assert "content of page 3" in out
    assert "content of page 1" not in out
    assert "content of page 5" not in out


def test_docx_sparse_breaks_fall_back_to_approx(tmp_path):
    """A long docx with only a FEW rendered page breaks must NOT be trusted as
    that few pages (Word writes them sporadically) — it falls back to char-approx
    so a high page number still resolves instead of selecting nothing.
    Regression: session-29 'summarize pages 68-70' returned empty."""
    from aiforge_core.runtime import doc_extract as de
    d = docx.Document()
    for i in range(1, 1400):                       # ~220k chars ≈ 70 printed pages
        d.add_paragraph(f"Para {i}. " + "word " * 30)
        if i % 140 == 0:                           # only ~10 manual breaks
            d.add_page_break()
    p = tmp_path / "long.docx"
    d.save(str(p))
    pages, kind = de.paginate(str(p), "")
    assert kind == "approx"                        # sparse breaks not trusted
    assert len(pages) > 60                         # spans the whole doc, not ~10
    out = de.extract_pages(str(p), "", "68-70")
    assert out
    assert "--- page 68 ---" in out
    assert "--- page 70 ---" in out


def _docx_with_reported_pages(docx, tmp_path, n_paras, reported):
    """Build a docx then inject Word's own <Pages> into docProps/app.xml."""
    import zipfile
    d = docx.Document()
    for i in range(1, n_paras):
        d.add_paragraph(f"Para {i}. " + "word " * 30)
    src = tmp_path / "src.docx"
    d.save(str(src))
    app = ('<?xml version="1.0"?><Properties xmlns="http://schemas.openxmlformats'
           f'.org/officeDocument/2006/extended-properties"><Pages>{reported}</Pages></Properties>')
    dst = tmp_path / "reported.docx"
    with zipfile.ZipFile(src) as zi, zipfile.ZipFile(dst, "w") as zo:
        for it in zi.namelist():
            zo.writestr(it, app if it == "docProps/app.xml" else zi.read(it))
    return str(dst)


def test_docx_honours_word_reported_page_count(tmp_path):
    """docx page count comes from Word's own docProps/app.xml <Pages> — the
    authoritative total. Regression: a 321-page doc was reported as 264."""
    from aiforge_core.runtime import doc_extract as de
    p = _docx_with_reported_pages(docx, tmp_path, 900, 321)
    assert de._docx_reported_pages(p) == 321
    pages, kind = de.paginate(p, "")
    assert len(pages) == 321          # EXACT match to Word's count
    assert kind == "word"
    out = de.extract_pages(p, "", "68-70")
    assert "--- page 68 ---" in out
    assert "--- page 70 ---" in out


def test_split_into_n_exact_count():
    from aiforge_core.runtime import doc_extract as de
    text = "\n".join(f"line {i} " + "x" * 40 for i in range(500))
    assert len(de._split_into_n(text, 321)) == 321      # many lines
    assert len(de._split_into_n("one\ntwo\nthree", 10)) == 10   # few lines
    assert de._split_into_n("solo", 1) == ["solo"]


def test_docx_dense_breaks_stay_exact(tmp_path):
    from aiforge_core.runtime import doc_extract as de
    d = docx.Document()
    for i in range(1, 40):
        d.add_paragraph(f"Page-{i} " + "word " * 200)   # ~1 page of prose each
        d.add_page_break()
    p = tmp_path / "dense.docx"
    d.save(str(p))
    _pages, kind = de.paginate(str(p), "")
    assert kind == "exact"


def test_summarize_doc_tool_out_of_range(monkeypatch, tmp_path):
    """The chat tool reports the real page count on an out-of-range request
    instead of returning empty (which made the model fumble)."""
    from aiforge_core.runtime.chat_agent._tools._code import _t_summarize_doc
    d = docx.Document()
    d.add_paragraph("only a little content here")
    media = tmp_path / ".aiforge" / "media"
    media.mkdir(parents=True)
    d.save(str(media / "small.docx"))
    r = _t_summarize_doc({"path": "small.docx", "pages": "68-70"}, str(tmp_path))
    assert r["ok"] is False
    assert "out of range" in r["error"]
    assert str(r["page_count"]) in r["error"]


def test_summarize_document_page_range(monkeypatch, tmp_path):
    """summarize_document(pages=...) extracts ONLY the range, then summarises."""
    from aiforge_core.runtime import doc_summarize as ds
    monkeypatch.setattr("aiforge_core.llm.client.complete", _fake_complete([]))
    monkeypatch.setenv("AIFORGE_SUMMARY_WINDOW_CHARS", "60")
    d = docx.Document()
    for i in range(1, 6):
        d.add_paragraph(f"detailed content of page {i} " + "z" * 80)
        if i < 5:
            d.add_page_break()
    p = tmp_path / "doc.docx"
    d.save(str(p))
    captured = {}
    orig = ds.summarize_text
    def _spy(text, role="chat"):
        captured["text"] = text
        return orig(text, role)
    monkeypatch.setattr(ds, "summarize_text", _spy)
    ds.summarize_document(str(p), role="chat", pages="2-3")
    assert "page 2" in captured["text"]
    assert "page 3" in captured["text"]
    assert "page 5" not in captured["text"]      # only the requested range


def test_describe_upload_summarizes_large_doc(monkeypatch, tmp_path):
    """Integration: a large text attachment gets a SUMMARY, a small one stays
    raw — the size auto-selection in chat_media._summarize_or_excerpt."""
    from aiforge_core.runtime import chat_media
    monkeypatch.setattr("aiforge_core.llm.client.complete",
                        _fake_complete([]))
    monkeypatch.setenv("AIFORGE_SUMMARY_WINDOW_CHARS", "500")
    big = tmp_path / "big.txt"
    big.write_text("\n".join(f"row {i} " + "q" * 90 for i in range(80)))
    out = chat_media.describe_upload(str(big), "big.txt", "text/plain", "chat")
    assert out.startswith("SUMMARY (auto-generated")

    small = tmp_path / "small.txt"
    small.write_text("just a line")
    out2 = chat_media.describe_upload(str(small), "small.txt", "text/plain", "chat")
    assert out2 == "just a line"

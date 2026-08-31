"""Turning an attached document into pages the model can be pointed at.

"Summarise pages 10-20 of the report" only works if page 10 here is page 10 in
the reader's viewer, and every format lies about pages differently:

  * PDF has real pages, and pypdf's flat text drops table structure — so
    pdfplumber's grid is folded back in when it is installed.
  * docx has no page concept at all. Word writes its OWN total into
    docProps/app.xml (authoritative), and sprinkles rendered page breaks
    sporadically — a 300-page report can carry ten markers, so trusting sparse
    markers under-counts massively and puts a real page number "out of range".
    Hence: Word's count first, dense breaks second, char estimate last, each
    labelled so the caller can warn that the numbers are approximate.
  * xlsx pages are sheets.

A page spec is clamped, not rejected: "10-99" on a 12-page file means pages
10..12, which is what someone asking for "10 onwards" meant.
"""
from __future__ import annotations

import types as pytypes
import zipfile

import pytest

from aiforge_core.runtime import doc_extract as D


# ─── knobs ─────────────────────────────────────────────────────────────


def test_a_junk_env_value_falls_back_to_the_default(monkeypatch):
    monkeypatch.setenv("AIFORGE_DOC_MAX_PAGES", "not-a-number")
    assert D._pdf_page_cap() == 500


def test_the_caps_are_tunable(monkeypatch):
    monkeypatch.setenv("AIFORGE_DOC_MAX_PAGES", "3")
    monkeypatch.setenv("AIFORGE_DOC_MAX_CHARS", "1234")
    assert D._pdf_page_cap() == 3 and D._doc_char_budget() == 1234


@pytest.mark.parametrize("path,mime,is_pdf", [
    ("a.pdf", "", True), ("a.PDF", "", True),
    ("a.bin", "application/pdf", True), ("a.docx", "", False)])
def test_a_pdf_is_recognised_by_either_name_or_mime(path, mime, is_pdf):
    assert D._is_pdf(path, mime) is is_pdf


@pytest.mark.parametrize("path,mime,is_docx", [
    ("a.docx", "", True),
    ("a.bin", "application/vnd.openxmlformats-officedocument.wordprocessingml"
     ".document", True),
    ("a.pdf", "", False)])
def test_a_docx_is_recognised_by_either_name_or_mime(path, mime, is_docx):
    assert D._is_docx(path, mime) is is_docx


# ─── PDF ───────────────────────────────────────────────────────────────


class _Page:
    def __init__(self, text="", tables=None, raise_on_text=False):
        self._text = text
        self._tables = tables or []
        self._raise = raise_on_text

    def extract_text(self):
        if self._raise:
            raise RuntimeError("encrypted")
        return self._text

    def extract_tables(self):
        return self._tables


@pytest.fixture()
def pdf(monkeypatch):
    """A stubbed pypdf + pdfplumber pair."""
    state: dict = {"pages": [_Page("page one"), _Page("page two")],
                   "tables": {}, "plumber_raises": False,
                   "no_plumber": False}
    import pypdf
    monkeypatch.setattr(pypdf, "PdfReader",
                        lambda path: pytypes.SimpleNamespace(
                            pages=state["pages"]))

    class _Plumbed:
        def __enter__(self):
            if state["plumber_raises"]:
                raise RuntimeError("bad pdf")
            return pytypes.SimpleNamespace(
                pages=[_Page(tables=state["tables"].get(i, []))
                       for i in range(len(state["pages"]))])

        def __exit__(self, *a):
            return False
    import pdfplumber
    monkeypatch.setattr(pdfplumber, "open", lambda path: _Plumbed())
    if state["no_plumber"]:
        pass
    return state


def test_each_pdf_page_is_its_own_entry(pdf, tmp_path):
    assert D._pdf_pages(str(tmp_path / "a.pdf")) == ["page one", "page two"]


def test_a_pages_tables_are_folded_back_in(pdf, tmp_path):
    """pypdf's flat text carries the numbers but loses the grid."""
    pdf["tables"] = {0: [[["Q1", "10"], ["Q2", None]]]}
    out = D._pdf_pages(str(tmp_path / "a.pdf"))
    assert "[tables]" in out[0] and "Q1 | 10" in out[0]
    assert out[0].endswith("Q2 |"), "an empty cell stays an empty cell"
    assert "[tables]" not in out[1]


def test_an_unreadable_page_does_not_lose_the_rest(pdf, tmp_path):
    pdf["pages"] = [_Page(raise_on_text=True), _Page("page two")]
    assert D._pdf_pages(str(tmp_path / "a.pdf")) == ["", "page two"]


def test_the_page_cap_bounds_a_huge_pdf(pdf, tmp_path, monkeypatch):
    monkeypatch.setenv("AIFORGE_DOC_MAX_PAGES", "1")
    pdf["pages"] = [_Page("one"), _Page("two"), _Page("three")]
    assert D._pdf_pages(str(tmp_path / "a.pdf")) == ["one"]


def test_a_pdf_that_pdfplumber_chokes_on_still_yields_its_text(pdf, tmp_path):
    pdf["plumber_raises"] = True
    assert D._pdf_pages(str(tmp_path / "a.pdf")) == ["page one", "page two"]


def test_without_pdfplumber_there_are_simply_no_tables(monkeypatch, tmp_path):
    import builtins
    real = builtins.__import__

    def _imp(name, *a, **k):
        if name == "pdfplumber":
            raise ImportError("not installed")
        return real(name, *a, **k)
    monkeypatch.setattr(builtins, "__import__", _imp)
    assert D._pdf_tables(str(tmp_path / "a.pdf"), 10) == {}


def test_a_tables_rows_are_flattened_with_newlines_squashed():
    page = _Page(tables=[[["a\nb", "c"]]])
    assert D._table_rows(page) == ["a b | c"]


# ─── docx ──────────────────────────────────────────────────────────────


def _docx(tmp_path, paragraphs, name="doc.docx", pages_reported=None):
    import docx
    d = docx.Document()
    for p in paragraphs:
        if p == "\f":
            d.add_page_break()
        else:
            d.add_paragraph(p)
    p = tmp_path / name
    d.save(str(p))
    if pages_reported is not None:
        _set_reported_pages(p, pages_reported)
    return str(p)


def _set_reported_pages(path, n):
    """Rewrite docProps/app.xml so Word's own <Pages> says ``n``."""
    import shutil
    tmp = str(path) + ".tmp"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/app.xml":
                data = (b'<?xml version="1.0"?><Properties xmlns="http://x">'
                        + f"<Pages>{n}</Pages>".encode() + b"</Properties>")
            zout.writestr(item, data)
    shutil.move(tmp, path)


def test_a_manual_page_break_starts_a_new_page(tmp_path):
    path = _docx(tmp_path, ["first page", "\f", "second page"])
    pages = D._docx_pages(path)
    assert pages[0] == "first page" and "second page" in pages[-1]


def test_a_table_lands_on_the_page_it_belongs_to(tmp_path):
    import docx
    d = docx.Document()
    d.add_paragraph("intro")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Q1"
    t.rows[0].cells[1].text = "10"
    p = tmp_path / "t.docx"
    d.save(str(p))
    assert "Q1 | 10" in D._docx_pages(str(p))[0]


def test_words_own_page_count_is_authoritative(tmp_path):
    """A 300-page report can carry ten break markers, so the markers alone
    under-count massively."""
    path = _docx(tmp_path, ["para " + "x" * 200] * 10, pages_reported=4)
    pages, kind = D._docx_paginate(path)
    assert kind == "word" and len(pages) == 4


def test_the_word_count_is_read_out_of_the_docx_zip(tmp_path):
    path = _docx(tmp_path, ["a"], pages_reported=7)
    assert D._docx_reported_pages(path) == 7


@pytest.mark.parametrize("value", ["0", "not-a-number", ""])
def test_an_unusable_word_count_is_ignored(tmp_path, value):
    path = _docx(tmp_path, ["a"])
    _set_reported_pages(path, value)
    assert D._docx_reported_pages(path) is None


def test_a_file_that_is_not_a_docx_reports_no_word_count(tmp_path):
    p = tmp_path / "plain.txt"
    p.write_text("hello")
    assert D._docx_reported_pages(str(p)) is None


def test_a_docx_without_a_pages_element_reports_nothing(tmp_path):
    path = _docx(tmp_path, ["a"])
    import shutil
    tmp = str(path) + ".t"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "docProps/app.xml":
                data = b'<?xml version="1.0"?><Properties xmlns="http://x"/>'
            zout.writestr(item, data)
    shutil.move(tmp, path)
    assert D._docx_reported_pages(path) is None


def test_dense_page_breaks_are_trusted_as_real_pages(tmp_path, monkeypatch):
    monkeypatch.setenv("AIFORGE_APPROX_PAGE_CHARS", "40")
    path = _docx(tmp_path, ["a" * 30, "\f", "b" * 30, "\f", "c" * 30])
    pages, kind = D._docx_paginate(path)
    assert kind == "exact" and len(pages) >= 3


def test_a_break_less_document_is_estimated(tmp_path, monkeypatch):
    monkeypatch.setenv("AIFORGE_APPROX_PAGE_CHARS", "50")
    path = _docx(tmp_path, ["x" * 40] * 6)
    pages, kind = D._docx_paginate(path)
    assert kind == "approx" and len(pages) > 1


def test_a_short_document_is_a_single_page(tmp_path):
    path = _docx(tmp_path, ["just a note"])
    assert D._docx_paginate(path) == (["just a note"], "approx")


def test_an_empty_document_has_no_pages(tmp_path):
    assert D._docx_paginate(_docx(tmp_path, [])) == ([], "none")


# ─── splitting to an exact page count ──────────────────────────────────


def test_text_is_split_into_exactly_the_count_word_reported():
    pages = D._split_into_n("\n".join(f"line {i}" for i in range(50)), 5)
    assert len(pages) == 5 and all(p.strip() for p in pages)


def test_a_document_with_fewer_lines_than_pages_still_fills_them():
    """Otherwise the tail pages come back empty and a real page number reads
    as blank."""
    pages = D._split_into_n("one long single line " * 20, 4)
    assert len(pages) == 4 and all(p.strip() for p in pages)


def test_a_single_page_split_is_the_whole_text():
    assert D._split_into_n("abc", 1) == ["abc"]


# ─── xlsx ──────────────────────────────────────────────────────────────


def _xlsx(tmp_path, sheets):
    from openpyxl import Workbook
    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for r in rows:
            ws.append(r)
    p = tmp_path / "book.xlsx"
    wb.save(str(p))
    return str(p)


def test_each_sheet_is_a_page(tmp_path):
    path = _xlsx(tmp_path, {"Sales": [["Q1", 10]], "Costs": [["Q1", 3]]})
    pages, kind = D.paginate(path)
    assert kind == "sheets" and len(pages) == 2
    assert pages[0].startswith("# sheet: Sales") and "Q1, 10" in pages[0]


def test_a_huge_sheet_is_cut_off_with_a_marker(tmp_path):
    path = _xlsx(tmp_path, {"Big": [[i] for i in range(600)]})
    assert "… (more rows)" in D.paginate(path)[0][0]


def test_an_empty_cell_is_left_empty(tmp_path):
    path = _xlsx(tmp_path, {"S": [["a", None, "c"]]})
    assert "a, , c" in D.paginate(path)[0][0]


# ─── plain text ────────────────────────────────────────────────────────


def test_a_small_text_file_is_one_exact_page(tmp_path):
    p = tmp_path / "notes.md"
    p.write_text("hello")
    assert D.paginate(str(p)) == (["hello"], "exact")


def test_a_large_text_file_is_estimated(tmp_path, monkeypatch):
    monkeypatch.setenv("AIFORGE_APPROX_PAGE_CHARS", "50")
    p = tmp_path / "notes.md"
    p.write_text("\n".join("x" * 40 for _ in range(10)))
    pages, kind = D.paginate(str(p))
    assert kind == "approx" and len(pages) > 1


def test_an_empty_text_file_has_nothing_to_read(tmp_path):
    p = tmp_path / "empty.txt"
    p.write_text("")
    assert D.paginate(str(p)) == ([], "none")


def test_a_format_nobody_handles_is_not_an_error(tmp_path):
    p = tmp_path / "image.png"
    p.write_bytes(b"\x89PNG")
    assert D.paginate(str(p)) == ([], "none")


def test_a_corrupt_document_never_raises(tmp_path):
    p = tmp_path / "broken.docx"
    p.write_bytes(b"not a zip")
    assert D.paginate(str(p)) == ([], "none")
    assert D.page_count(str(p)) == 0
    assert D.extract_text(str(p)) == ""


def test_the_pagination_kind_is_reportable(tmp_path):
    p = tmp_path / "n.txt"
    p.write_text("hi")
    assert D.pagination_kind(str(p)) == "exact"
    assert D.document_pages(str(p)) == ["hi"]


# ─── page markers and the char budget ──────────────────────────────────


def test_every_page_is_labelled_so_the_model_can_cite_it():
    out = D._join_pages(["one", "two"], budget=10_000)
    assert out.startswith("--- page 1 ---\none") and "--- page 2 ---" in out


def test_the_labels_follow_the_pages_actually_selected():
    out = D._join_pages(["ten", "eleven"], budget=10_000, numbers=[10, 11])
    assert "--- page 10 ---" in out and "--- page 11 ---" in out


def test_the_join_stops_at_the_budget_and_says_where():
    out = D._join_pages(["x" * 100, "y" * 100, "z" * 100], budget=120)
    assert "truncated at page 2" in out and "z" * 100 not in out


def test_the_whole_text_of_a_document_carries_its_page_markers(tmp_path):
    path = _docx(tmp_path, ["first", "\f", "second"])
    assert "--- page 1 ---" in D.extract_text(path)


def test_a_text_file_is_returned_as_it_is(tmp_path):
    p = tmp_path / "n.txt"
    p.write_text("raw text")
    assert D.extract_text(str(p)) == "raw text"


def test_an_unreadable_file_extracts_to_nothing(tmp_path):
    assert D.extract_text(str(tmp_path / "ghost.txt")) == ""


# ─── page selection ────────────────────────────────────────────────────


@pytest.mark.parametrize("spec,expected", [
    ("1", [0]),
    ("3,5", [2, 4]),
    ("2-4", [1, 2, 3]),
    ("4-2", [1, 2, 3]),          # backwards range is normalised
    ("1, 2 , 3", [0, 1, 2]),     # spaces are ignored
    ("2,2,2", [1]),              # duplicates collapse
])
def test_a_page_spec_becomes_indices(spec, expected):
    assert D.parse_page_spec(spec, 10) == expected


def test_a_range_past_the_end_is_clamped_not_rejected():
    """"10-99" on a 12-page file means 10..12 — what the reader meant."""
    assert D.parse_page_spec("10-99", 12) == [9, 10, 11]


@pytest.mark.parametrize("spec", ["", "abc", "1-x", "-", None])
def test_garbage_selects_nothing(spec):
    assert D.parse_page_spec(spec, 10) == []


def test_a_page_before_the_first_is_dropped():
    assert D.parse_page_spec("0-2", 10) == [0, 1]


def test_a_spec_against_an_empty_document_selects_nothing():
    assert D.parse_page_spec("1", 0) == []


def test_only_the_asked_for_pages_come_back(tmp_path):
    path = _xlsx(tmp_path, {"A": [["a"]], "B": [["b"]], "C": [["c"]]})
    out = D.extract_pages(path, "", "2-3")
    assert "--- page 2 ---" in out and "sheet: B" in out
    assert "sheet: A" not in out


def test_selecting_nothing_returns_nothing(tmp_path):
    path = _xlsx(tmp_path, {"A": [["a"]]})
    assert D.extract_pages(path, "", "9") == ""


def test_the_page_count_is_what_this_module_segments(tmp_path):
    path = _xlsx(tmp_path, {"A": [["a"]], "B": [["b"]]})
    assert D.page_count(path) == 2
